# legal/views.py
from rest_framework import viewsets, status
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from django.http import HttpResponse
from .models import PowerOfAttorney
from .serializers import PowerOfAttorneySerializer, PowerOfAttorneyListSerializer
from docx import Document
import os
from io import BytesIO
from datetime import datetime
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from django.conf import settings

class PowerOfAttorneyViewSet(viewsets.ModelViewSet):
    """
    ViewSet for handling CRUD operations for Power of Attorney documents
    and document generation in Word or PDF format.
    """
    queryset = PowerOfAttorney.objects.all()
    permission_classes = [IsAuthenticated]
    lookup_field = 'slug'

    def get_serializer_class(self):
        if self.action == 'list':
            return PowerOfAttorneyListSerializer
        return PowerOfAttorneySerializer

    def get_queryset(self):
        return super().get_queryset()

    def perform_create(self, serializer):
        serializer.save()

    def perform_update(self, serializer):
        serializer.save()

    @action(detail=True, methods=['get'], url_path='generate-document')
    def generate_document(self, request, slug=None):
        """
        Generate document in specified format (docx or pdf).
        Query parameter 'format' can be 'docx' or 'pdf'.
        """
        instance = self.get_object()
        format_type = request.query_params.get('format', 'docx').lower()

        if format_type not in ['docx', 'pdf']:
            return Response(
                {'error': "Format must be 'docx' or 'pdf'"},
                status=status.HTTP_400_BAD_REQUEST
            )

        try:
            if format_type == 'docx':
                return self._generate_word_document(instance)
            return self._generate_pdf_document(instance)
        except Exception as e:
            return Response(
                {'error': f"Error generating document: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    def _generate_word_document(self, instance):
        template_path = os.path.join(settings.BASE_DIR, 'templates', 'power_of_attorney_template.docx')
        if not os.path.exists(template_path):
            raise Exception("Document template not found")

        doc = Document(template_path)
        replacements = {
            '{DOCUMENT_DATE}': instance.document_date.strftime('%dth, %B %Y'),
            '{COMPANY_NAME}': instance.company_name,
            '{COMPANY_ADDRESS}': instance.company_address,
            '{COMPANY_PO_BOX}': instance.company_po_box,
            '{ATTORNEY_NAME}': instance.attorney_name,
            '{ATTORNEY_PO_BOX}': instance.attorney_po_box,
            '{ATTORNEY_ADDRESS}': instance.attorney_address or '',
            '{TENDER_NUMBER}': instance.tender_number,
            '{TENDER_DESCRIPTION}': instance.tender_description,
            '{TENDER_BENEFICIARY}': instance.tender_beneficiary,
            '{WITNESS_NAME}': instance.witness_name,
            '{WITNESS_PO_BOX}': instance.witness_po_box,
            '{WITNESS_TITLE}': instance.witness_title,
            '{WITNESS_ADDRESS}': instance.witness_address or ''
        }

        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    for run in paragraph.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, value)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                for run in paragraph.runs:
                                    if key in run.text:
                                        run.text = run.text.replace(key, value)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename=poa_{instance.slug}.docx'
        return response

    def _generate_pdf_document(self, instance):
        buffer = BytesIO()
        p = canvas.Canvas(buffer, pagesize=letter)
        
        p.setFont("Helvetica-Bold", 14)
        p.drawString(100, 750, "STANDARD POWER OF ATTORNEY")
        p.setFont("Helvetica", 12)
        p.drawString(100, 730, "TO ALL IT MAY CONCERN")

        content = f"""
        THAT BY THIS POWER OF ATTORNEY given on the {instance.document_date.strftime('%dth, %B %Y')}, 
        WE the undersigned {instance.company_name} of {instance.company_address}, {instance.company_po_box}, 
        do hereby ordain, nominate, authorize, empower and appoint {instance.attorney_name} of 
        {instance.attorney_po_box}{', ' + instance.attorney_address if instance.attorney_address else ''}, 
        to be our true lawful Attorney and Agent, with full power and authority, for us and in our names, 
        and for our accounts and benefits, to do any, or all of the following acts, in the execution of tender 
        No. {instance.tender_number} for {instance.tender_description} for the {instance.tender_beneficiary};
        
        To act for the company and do any other thing or things incidental for {instance.tender_number} 
        for {instance.tender_description};
        
        BEFORE ME;
        {instance.witness_name}, {instance.witness_title}
        {instance.witness_po_box}
        {instance.witness_address if instance.witness_address else ''}
        """

        y_position = 700
        for line in content.split('\n'):
            line = line.strip()
            if line:
                p.drawString(100, y_position, line)
                y_position -= 15
                if y_position < 50:
                    p.showPage()
                    p.setFont("Helvetica", 12)
                    y_position = 750

        p.showPage()
        p.save()

        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/pdf'
        )
        response['Content-Disposition'] = f'attachment; filename=poa_{instance.slug}.pdf'
        buffer.close()
        return response